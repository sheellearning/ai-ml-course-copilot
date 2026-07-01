import os
from typing import List, Dict, Any

from docx import Document


# Load text from DOCX paragraphs and tables with source metadata.
def load_docx_file(docx_path: str) -> List[Dict[str, Any]]:
    doc = Document(docx_path)
    records = []
    filename = os.path.basename(docx_path)

    paragraphs = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    # Preserve table rows as pipe-separated text.
    for table in doc.tables:
        for row in table.rows:
            row_text = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                paragraphs.append(" | ".join(row_text))

    full_text = "\n".join(paragraphs).strip()

    if not full_text:
        return records

    metadata = {
        "source": "course",
        "doc_type": "docx",
        "filename": filename,
        "path": docx_path,
    }

    records.append(
        {
            "text": full_text,
            "metadata": metadata,
        }
    )

    return records