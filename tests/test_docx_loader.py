from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from core.chunking import chunk_records
from core.loaders.docx_loader import load_docx_file


# Create a temporary DOCX containing paragraphs and a table.
def create_sample_docx(docx_path):
    document = Document()
    document.add_paragraph("Transformer models process sequential data.")

    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Model"
    table.rows[0].cells[1].text = "BERT"

    document.save(docx_path)


# Verify DOCX text extraction, metadata, and chunk creation.
def test_docx_loader():
    with TemporaryDirectory() as temporary_folder:
        docx_path = Path(temporary_folder) / "sample.docx"
        create_sample_docx(docx_path)

        records = load_docx_file(str(docx_path))
        chunks = chunk_records(records)

        assert len(records) == 1
        assert chunks

        record = records[0]

        assert "Transformer models process sequential data." in record["text"]
        assert "Model | BERT" in record["text"]
        assert record["metadata"]["source"] == "course"
        assert record["metadata"]["doc_type"] == "docx"
        assert record["metadata"]["filename"] == "sample.docx"
        assert record["metadata"]["path"] == str(docx_path)

        assert chunks[0]["metadata"]["chunk_strategy"] == "recursive_docx"


if __name__ == "__main__":
    test_docx_loader()
    print("DOCX LOADER TEST PASSED")